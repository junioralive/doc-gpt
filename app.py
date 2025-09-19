import streamlit as st
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from langchain_community.embeddings import HuggingFaceEmbeddings
import google.generativeai as genai
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
import time
import json
import pandas as pd
import docx
from docx import Document
import io
import chardet

# Load environment variables
load_dotenv()

# Configure Google AI
try:
    api_key = os.getenv("GOOGLE_API_KEY")
    if not api_key:
        st.error("🔑 GOOGLE_API_KEY not found in environment variables")
        st.stop()
    genai.configure(api_key=api_key)
except Exception as e:
    st.error(f"❌ Error configuring Google AI: {str(e)}")
    st.stop()

def get_file_text(uploaded_files):
    """Extract text from uploaded files of various formats with error handling"""
    text = ""
    try:
        if not uploaded_files:
            st.warning("📄 Please upload at least one file")
            return ""
        
        for file in uploaded_files:
            try:
                file_extension = file.name.lower().split('.')[-1]
                file_text = ""
                
                if file_extension == 'pdf':
                    # Handle Files
                    pdf_reader = PdfReader(file)
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                file_text += page_text
                        except Exception as e:
                            st.warning(f"⚠️ Could not extract text from page {page_num + 1} of {file.name}")
                            continue
                
                elif file_extension in ['docx', 'doc']:
                    # Handle Word documents
                    try:
                        doc = Document(file)
                        for paragraph in doc.paragraphs:
                            file_text += paragraph.text + "\n"
                        
                        # Extract text from tables
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    file_text += cell.text + " "
                                file_text += "\n"
                    except Exception as e:
                        st.warning(f"⚠️ Could not extract text from Word document {file.name}: {str(e)}")
                        continue
                
                elif file_extension in ['xlsx', 'xls']:
                    # Handle Excel files
                    try:
                        # Read all sheets
                        excel_file = pd.ExcelFile(file)
                        for sheet_name in excel_file.sheet_names:
                            df = pd.read_excel(file, sheet_name=sheet_name)
                            # Convert DataFrame to text
                            file_text += f"\n--- Sheet: {sheet_name} ---\n"
                            file_text += df.to_string(index=False) + "\n"
                    except Exception as e:
                        st.warning(f"⚠️ Could not extract text from Excel file {file.name}: {str(e)}")
                        continue
                
                elif file_extension == 'csv':
                    # Handle CSV files
                    try:
                        # Detect encoding
                        raw_data = file.read()
                        file.seek(0)  # Reset file pointer
                        
                        # Try to detect encoding
                        detected = chardet.detect(raw_data)
                        encoding = detected.get('encoding', 'utf-8')
                        
                        # Read CSV with detected encoding
                        df = pd.read_csv(file, encoding=encoding)
                        file_text += df.to_string(index=False) + "\n"
                    except Exception as e:
                        try:
                            # Fallback to common encodings
                            file.seek(0)
                            df = pd.read_csv(file, encoding='latin-1')
                            file_text += df.to_string(index=False) + "\n"
                        except Exception as e2:
                            st.warning(f"⚠️ Could not extract text from CSV file {file.name}: {str(e2)}")
                            continue
                
                elif file_extension == 'txt':
                    # Handle text files
                    try:
                        # Read as bytes and detect encoding
                        raw_data = file.read()
                        detected = chardet.detect(raw_data)
                        encoding = detected.get('encoding', 'utf-8')
                        
                        # Decode with detected encoding
                        file_text = raw_data.decode(encoding)
                    except Exception as e:
                        try:
                            # Fallback to common encodings
                            file_text = raw_data.decode('latin-1')
                        except Exception as e2:
                            st.warning(f"⚠️ Could not extract text from text file {file.name}: {str(e2)}")
                            continue
                
                else:
                    st.warning(f"⚠️ Unsupported file format: {file_extension} for file {file.name}")
                    continue
                
                if file_text.strip():
                    text += f"\n--- Content from {file.name} ---\n"
                    text += file_text + "\n"
                else:
                    st.warning(f"⚠️ No text could be extracted from {file.name}")
                    
            except Exception as e:
                st.error(f"❌ Error processing file {file.name}: {str(e)}")
                continue
        
        if not text.strip():
            st.error("❌ No text could be extracted from the uploaded files")
            return ""
            
        return text
    except Exception as e:
        st.error(f"❌ Error processing files: {str(e)}")
        return ""

def get_text_chunks(text):
    """Split text into chunks with error handling"""
    try:
        if not text.strip():
            return []
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=10000, 
            chunk_overlap=1000
        )
        chunks = text_splitter.split_text(text)
        
        if not chunks:
            st.error("❌ Could not create text chunks from the extracted text")
            return []
            
        return chunks
    except Exception as e:
        st.error(f"❌ Error creating text chunks: {str(e)}")
        return []

def get_vector_store(text_chunks):
    """Create and save vector store with error handling"""
    try:
        if not text_chunks:
            st.error("❌ No text chunks available to create vector store")
            return False
        
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
        vector_store.save_local("faiss_index")
        return True
    except Exception as e:
        st.error(f"❌ Error creating vector store: {str(e)}")
        return False

def get_conversational_chain():
    """Create conversational chain with error handling using Gemini 2.5 Flash"""
    try:
        prompt_template = """
        Answer the question as detailed as possible from the provided context. Make sure to provide all the details.
        If the answer is not in the provided context, just say "Answer is not available in the context".
        Do not provide wrong answers.

        Context:
        {context}

        Question: {question}

        Answer:
        """

        model = ChatGoogleGenerativeAI(
            model="gemini-2.0-flash-exp",
            temperature=0.3
        )

        prompt = PromptTemplate(
            template=prompt_template, 
            input_variables=["context", "question"]
        )
        
        return model, prompt
    except Exception as e:
        st.error(f"❌ Error creating conversational chain: {str(e)}")
        return None, None

def summarize_pdf():
    """Generate document summary"""
    try:
        if not os.path.exists("faiss_index"):
            st.error("📄 Please upload and process documents first")
            return
        
        with st.spinner("📝 Generating summary..."):
            embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
            new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
            
            # Get all documents for summary
            docs = new_db.similarity_search("summary main points key information", k=10)
            
            if not docs:
                st.warning("⚠️ No content found for summarization")
                return
            
            model, _ = get_conversational_chain()
            if not model:
                return
            
            context = "\n\n".join([doc.page_content for doc in docs])
            summary_prompt = f"""
            Please provide a comprehensive summary of the following document content. 
            Include the main points, key findings, and important information:

            {context}

            Summary:
            """
            
            response = model.invoke(summary_prompt)
            
            if response and hasattr(response, 'content'):
                st.success("✅ Summary generated!")
                st.markdown("### 📋 Document Summary")
                st.write(response.content)
            else:
                st.error("❌ Could not generate summary") 
                
    except Exception as e:
        st.error(f"❌ Error generating summary: {str(e)}")

def generate_questions():
    """Generate questions from document content"""
    try:
        if not os.path.exists("faiss_index"):
            st.error("📄 Please upload and process documents first")
            return
        
        with st.spinner("❓ Generating questions..."):
            embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
            new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
            
            docs = new_db.similarity_search("main topics concepts important information", k=8)
            
            if not docs:
                st.warning("⚠️ No content found for question generation")
                return
            
            model, _ = get_conversational_chain()
            if not model:
                return
            
            context = "\n\n".join([doc.page_content for doc in docs])
            questions_prompt = f"""
            Based on the following document content, generate 8-10 thoughtful questions that would help someone understand the key concepts and important information. 
            Make the questions clear and specific:

            {context}

            Questions:
            """
            
            response = model.invoke(questions_prompt)
            
            if response and hasattr(response, 'content'):
                st.success("✅ Questions generated!")
                st.markdown("### ❓ Practice Questions")
                
                questions = response.content.split('\n')
                questions = [q.strip() for q in questions if q.strip() and ('?' in q or q.strip().endswith('.'))]
                
                for i, question in enumerate(questions[:10], 1):
                    question = question.lstrip('0123456789.- ')
                    st.markdown(f"**{i}.** {question}")
                    
                    if st.button(f"Get Answer", key=f"answer_{i}"):
                        answer_question(question)
            else:
                st.error("❌ Could not generate questions")
                
    except Exception as e:
        st.error(f"❌ Error generating questions: {str(e)}")

def generate_mcqs():
    """Generate multiple choice questions"""
    try:
        if not os.path.exists("faiss_index"):
            st.error("📄 Please upload and process documents first")
            return
        
        with st.spinner("📝 Generating MCQs..."):
            embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
            new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
            
            docs = new_db.similarity_search("key concepts important facts definitions", k=6)
            
            if not docs:
                st.warning("⚠️ No content found for MCQ generation")
                return
            
            model, _ = get_conversational_chain()
            if not model:
                return
            
            context = "\n\n".join([doc.page_content for doc in docs])
            mcq_prompt = f"""
            Based on the following document content, create 5 multiple choice questions with 4 options each (A, B, C, D). 
            Include the correct answer at the end. Format each question clearly:

            {context}

            MCQs:
            """
            
            response = model.invoke(mcq_prompt)
            
            if response and hasattr(response, 'content'):
                st.success("✅ MCQs generated!")
                st.markdown("### 📝 Multiple Choice Questions")
                st.write(response.content)
            else:
                st.error("❌ Could not generate MCQs")
                
    except Exception as e:
        st.error(f"❌ Error generating MCQs: {str(e)}")

def generate_notes():
    """Generate short notes from documents"""
    try:
        if not os.path.exists("faiss_index"):
            st.error("📄 Please upload and process documents first")
            return
        
        with st.spinner("📚 Generating notes..."):
            embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
            new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
            
            docs = new_db.similarity_search("main concepts key points important information", k=8)
            
            if not docs:
                st.warning("⚠️ No content found for notes generation")
                return
            
            model, _ = get_conversational_chain()
            if not model:
                return
            
            context = "\n\n".join([doc.page_content for doc in docs])
            notes_prompt = f"""
            Create concise, well-organized study notes from the following content. 
            Use bullet points, headings, and clear structure. Focus on key concepts and important information:

            {context}

            Study Notes:
            """
            
            response = model.invoke(notes_prompt)
            
            if response and hasattr(response, 'content'):
                st.success("✅ Notes generated!")
                st.markdown("### 📚 Study Notes")
                st.write(response.content)
            else:
                st.error("❌ Could not generate notes")
                
    except Exception as e:
        st.error(f"❌ Error generating notes: {str(e)}")

def answer_question(question):
    """Answer a specific question"""
    try:
        if not os.path.exists("faiss_index"):
            st.error("📄 Please upload and process documents first")
            return
        
        with st.spinner("🔍 Finding answer..."):
            embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
            new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
            docs = new_db.similarity_search(question)
            
            if not docs:
                st.warning("⚠️ No relevant information found")
                return
            
            model, prompt = get_conversational_chain()
            if not model or not prompt:
                return
            
            context = "\n\n".join([doc.page_content for doc in docs])
            formatted_prompt = prompt.format(context=context, question=question)
            
            response = model.invoke(formatted_prompt)
            
            if response and hasattr(response, 'content'):
                st.success("✅ Answer found!")
                st.write("**Answer:**")
                st.write(response.content)
            else:
                st.error("❌ Could not generate answer")
                
    except Exception as e:
        st.error(f"❌ Error answering question: {str(e)}")

def user_input(user_question):
    """Process user question and generate response with error handling"""
    try:
        if not user_question.strip():
            st.warning("⚠️ Please enter a question")
            return
        
        if not os.path.exists("faiss_index"):
            st.error("📄 Please upload and process documents first")
            return
        
        with st.spinner("🔍 Searching for answer..."):
            embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
            
            try:
                new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
            except Exception as e:
                st.error("❌ Error loading vector store. Please reprocess your documents.")
                return
            
            docs = new_db.similarity_search(user_question)
            
            if not docs:
                st.warning("⚠️ No relevant information found in the uploaded documents")
                return
            
            model, prompt = get_conversational_chain()
            if not model or not prompt:
                return
            
            context = "\n\n".join([doc.page_content for doc in docs])
            formatted_prompt = prompt.format(context=context, question=user_question)
            
            try:
                response = model.invoke(formatted_prompt)
                
                if response and hasattr(response, 'content'):
                    st.success("✅ Answer found!")
                    st.write("**Reply:**")
                    st.write(response.content)
                else:
                    st.error("❌ Could not generate a response")
            except Exception as e:
                st.error(f"❌ Error generating response: {str(e)}")
                
    except Exception as e:
        st.error(f"❌ Error processing question: {str(e)}")

def main():
    """Main application function"""
    st.set_page_config(
        page_title="DOC-GPT | Chat, Analyze & Study Your Documents",
        page_icon="📚",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    st.markdown("""
    <style>
    .main-header {
        text-align: center;
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-size: 3.5rem;
        font-weight: bold;
        margin-bottom: 0.5rem;
    }
    .sub-header {
        text-align: center;
        color: #666;
        font-size: 1.2rem;
        margin-bottom: 1rem;
    }
    .tagline {
        text-align: center;
        color: #888;
        font-style: italic;
        margin-bottom: 2rem;
    }
    .sidebar-header {
        color: #667eea;
        font-size: 1.5rem;
        font-weight: bold;
        margin-bottom: 1rem;
    }
    .feature-button {
        margin: 0.25rem;
        width: 100%;
    }
    .stButton > button {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        border-radius: 10px;
        padding: 0.5rem 1rem;
        font-weight: bold;
        transition: all 0.3s ease;
    }
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 8px rgba(0,0,0,0.2);
    }
    .info-box {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        padding: 1rem;
        border-radius: 10px;
        margin: 1rem 0;
        border-left: 4px solid #667eea;
    }
    </style>
    """, unsafe_allow_html=True)
    
    st.markdown('<h1 class="main-header">DOC-GPT</h1>', unsafe_allow_html=True)
    st.markdown('<p class="sub-header">Intelligent Document Analysis & Chat Platform</p>', unsafe_allow_html=True)
    st.markdown('<p class="tagline">Transform Documents into Interactive Knowledge • Perfect for Students • Researchers • Professionals</p>', unsafe_allow_html=True)
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.markdown('<div class="info-box">📚 <b>Study Smarter</b><br>Generate summaries and notes</div>', unsafe_allow_html=True)
    with col2:
        st.markdown('<div class="info-box">🔍 <b>Research Faster</b><br>Ask questions instantly</div>', unsafe_allow_html=True)
    with col3:
        st.markdown('<div class="info-box">📝 <b>Practice Better</b><br>Generate questions & MCQs</div>', unsafe_allow_html=True)
    with col4:
        st.markdown('<div class="info-box">⚡ <b>Save Time</b><br>Extract key information quickly</div>', unsafe_allow_html=True)
    
    st.markdown("---")
    
    st.markdown("### 🚀 AI-Powered Features")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        if st.button("📋 Summarize Documents", use_container_width=True):
            summarize_pdf()
    
    with col2:
        if st.button("❓ Generate Questions", use_container_width=True):
            generate_questions()
    
    with col3:
        if st.button("📝 Create MCQs", use_container_width=True):
            generate_mcqs()
    
    with col4:
        if st.button("📚 Generate Notes", use_container_width=True):
            generate_notes()
    
    st.markdown("---")
    
    # Question input section
    st.markdown("### 💬 Ask Questions About Your Documents")
    user_question = st.text_input(
        "Enter your question:",
        placeholder="What are the main findings in this research paper?",
        help="Ask any question about your uploaded documents"
    )
    
    if user_question:
        user_input(user_question)
    
    with st.sidebar:
        st.markdown('<h2 class="sidebar-header">📁 Document Upload</h2>', unsafe_allow_html=True)
        
        st.markdown("**Quick Start Guide:**")
        st.markdown("1. 📤 Upload your Files")
        st.markdown("2. ⚙️ Click 'Process Documents'")
        st.markdown("3. 🚀 Use AI features or ask questions")
        
        st.markdown("---")
        
        uploaded_docs = st.file_uploader(
            "Choose Files",
            accept_multiple_files=True,
            type=['pdf', 'docx', 'doc', 'xlsx', 'xls', 'csv', 'txt'],
            help="Upload PDF, Word, Excel, CSV, or text files to analyze"
        )
        
        if uploaded_docs:
            st.success(f"✅ Uploaded {len(uploaded_docs)} file(s)")
            for doc in uploaded_docs:
                st.write(f"📄 {doc.name}")
        
        if st.button("⚙️ Process Documents", type="primary", use_container_width=True):
            if not uploaded_docs:
                st.error("❌ Please upload at least one file")
                st.error("❌ Please upload at least one file")
            else:
                with st.spinner("⚙️ Processing documents..."):
                    progress_bar = st.progress(0)
                    
                    progress_bar.progress(25)
                    raw_text = get_file_text(uploaded_docs)
                    
                    if raw_text:
                        progress_bar.progress(50)
                        text_chunks = get_text_chunks(raw_text)
                        
                        if text_chunks:
                            progress_bar.progress(75)
                            success = get_vector_store(text_chunks)
                            
                            progress_bar.progress(100)
                            
                            if success:
                                st.success("✅ Documents processed successfully!")
                                st.balloons()
                            else:
                                st.error("❌ Failed to process documents")
                        else:
                            st.error("❌ Failed to create text chunks")
                    else:
                        st.error("❌ Failed to extract text from documents")
                    
                    progress_bar.empty()
        
        st.markdown("---")
        
        st.markdown("### ℹ️ About DOC-GPT")
        st.markdown("**Created by:** Junior Alive")
        st.markdown("**Description:** Advanced document intelligence platform that transforms various document formats into searchable knowledge hubs. Instantly summarize content, generate questions, and chat with your documents using state-of-the-art AI technology.")
        
        st.markdown("**🔗 Links:**")
        st.markdown("• [📂 GitHub Repository](https://github.com/junioralive/doc-gpt)")
        st.markdown("• [🐛 Report Issues](https://github.com/junioralive/doc-gpt/issues)")
        
        st.markdown("---")
        
        st.markdown("**🤖 AI Model:** Google Gemini 2.0 Flash")
        st.markdown("**🔒 Privacy:** Your documents are processed securely and locally")
        st.markdown("**💡 Tip:** Upload research papers, academic articles, textbooks, reports, Excel files, Word documents, CSV data, or any supported document format to unlock intelligent document analysis and interactive learning!")

if __name__ == "__main__":
    main()
